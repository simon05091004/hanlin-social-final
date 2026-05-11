from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import B4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

from pathlib import Path


ROOT = Path(__file__).resolve().parent
DOCX_OUT = ROOT / "翰林五下社會期末複習卷_B4.docx"
PDF_OUT = ROOT / "翰林五下社會期末複習卷_B4.pdf"


EXAM = {
    "title": "國小五下社會期末複習卷",
    "subtitle": "翰林版主軸｜B4 兩張雙面｜總分 100 分",
    "sections": [
        {
            "heading": "第1面：基礎概念與判斷題（28分）",
            "blocks": [
                ("h2", "一、選擇題：每題2分，共20分"),
                ("q", "1. 臺灣各區域發展不同，最可能和下列哪一組因素有關？\nA. 星座、血型、人口年齡　B. 地形、交通、歷史、產業\nC. 電視節目、運動項目、飲食習慣　D. 學校數量、班級人數、考試次數"),
                ("q", "2. 東部地區發展較晚，主要原因之一是：\nA. 完全沒有河川　B. 受到山脈阻隔，交通較不便利\nC. 沒有任何自然景觀　D. 人們不能從事觀光活動"),
                ("q", "3. 下列哪一項最能說明「交通建設影響區域發展」？\nA. 高速公路通車後，南北往來更方便　B. 學生每天要寫聯絡簿\nC. 天氣熱時大家穿短袖　D. 考試前要複習"),
                ("q", "4. 臺灣經濟發展歷程較合理的順序是：\nA. 高科技→農業→工業　B. 工業→農業→高科技\nC. 農業→工業→高科技　D. 農業→高科技→工業"),
                ("q", "5. 十大建設對臺灣的主要影響是：\nA. 讓臺灣完全沒有環境問題　B. 促進交通、能源、工業與都市發展\nC. 使所有人都搬到離島　D. 讓農業完全消失"),
                ("q", "6. 臺灣能發展高科技產業，和下列哪一項關係最密切？\nA. 科學園區、人才培育、研發投資　B. 只有靠天氣好\nC. 完全不需要教育　D. 只靠進口產品"),
                ("q", "7. 製作小書時，最適合的第一步是：\nA. 先畫裝飾圖　B. 先決定主題和想探究的問題　C. 先裝訂　D. 先寫心得"),
                ("q", "8. 查詢近期社會事件，最適合參考：\nA. 新聞或報紙　B. 十年前的日記　C. 猜測　D. 沒有來源的留言"),
                ("q", "9. 民主社會中，人民參與公共事務的方式不包括：\nA. 選舉　B. 理性討論　C. 關心公共議題　D. 用暴力強迫別人接受意見"),
                ("q", "10. 臺灣多元文化的正確態度是：\nA. 只接受自己的文化　B. 尊重、理解並欣賞不同文化\nC. 覺得不同文化都不重要　D. 要求所有人生活方式完全一樣"),
                ("h2", "二、配對題：每題2分，共8分"),
                ("table", [["11. 北部地區", "A. 自然景觀豐富，觀光資源多"], ["12. 東部地區", "B. 政治、經濟、交通中心"], ["13. 實地踏查", "C. 可以取得第一手資料"], ["14. 高科技產業", "D. 需要研發、人才與資金"]]),
            ],
        },
        {
            "heading": "第2面：圖表判讀與時序題（24分）",
            "blocks": [
                ("h2", "三、圖表判讀題：每題4分，共16分"),
                ("p", "請閱讀下表回答問題。"),
                ("table", [["時期", "主要產業特色", "可能影響"], ["早期", "農業為主", "人口多分布在適合耕作地區"], ["後來", "工業發展", "工廠增加，都市人口成長"], ["現代", "高科技產業發展", "需要專業人才與研發能力"]]),
                ("q", "15. 從表中可看出臺灣產業大致如何轉變？"),
                ("q", "16. 工業發展後，為什麼都市人口可能增加？"),
                ("q", "17. 現代高科技產業需要哪些條件？請寫出兩項。"),
                ("q", "18. 經濟發展可能帶來便利，也可能造成什麼問題？請寫出一項。"),
                ("h2", "四、時序與因果題：共8分"),
                ("q", "19. 請將下列事件依時間先後排列，填入正確順序。每格1分，共4分。\nA. 解嚴後民主發展更進一步　B. 日本統治臺灣\nC. 戰後中華民國政府治理臺灣　D. 臺灣進入高科技產業發展階段\n正確順序：＿＿ → ＿＿ → ＿＿ → ＿＿"),
                ("q", "20. 請從上題選一個事件，說明它對臺灣社會的影響。4分。"),
            ],
        },
        {
            "heading": "第3面：情境素養題（28分）",
            "blocks": [
                ("h2", "五、生活情境題：每題7分，共28分"),
                ("q", "21. 小安一家從花蓮到臺北看病，爸爸說：「現在交通方便很多，但東部到西部還是會受地形影響。」\n（1）東部交通較受限制，和哪種自然因素有關？2分\n（2）交通改善會對生活造成哪些影響？寫兩點。4分\n（3）請判斷爸爸的說法是否合理，並說明原因。1分"),
                ("q", "22. 某地興建工業區後，就業機會增加，但附近居民也擔心空氣與水汙染。\n（1）工業區可能帶來什麼好處？2分\n（2）可能帶來什麼問題？2分\n（3）如果你是地方居民，你會提出什麼建議，讓發展和環境取得平衡？3分"),
                ("q", "23. 小組要做一本「臺灣茶文化」小書，有人找到網路文章，有人訪問茶農，有人拍攝茶園照片。\n（1）訪問茶農屬於哪一種資料蒐集方式？2分\n（2）網路資料使用前要注意什麼？2分\n（3）小書呈現時，怎樣做比較容易讓讀者看懂？寫兩點。3分"),
                ("q", "24. 班上討論是否減少一次性餐具。有同學說：「我覺得麻煩，所以不用管環保。」另一位同學說：「我們可以先從班級活動少用免洗餐具開始。」\n（1）哪一位同學比較符合公民參與精神？1分\n（2）為什麼？3分\n（3）請提出一個你可做到的具體行動。3分"),
            ],
        },
        {
            "heading": "第4面：閱讀理解、簡答與答案解析（20分）",
            "blocks": [
                ("h2", "六、資料閱讀題：每題4分，共12分"),
                ("p", "臺灣不同區域有不同特色。北部交通便利、工商業發達；中部位居南北交通要道，農業與工業都有發展；南部開發較早，農業、港口與工業發展重要；東部自然景觀豐富，但因山脈阻隔，交通發展較不容易；離島則受到海洋交通影響，生活和產業常需要與本島連結。"),
                ("q", "25. 根據資料，北部的重要特色是什麼？"),
                ("q", "26. 東部發展受到哪一項自然因素影響？"),
                ("q", "27. 如果要介紹離島生活，為什麼不能只看地圖，還要了解交通與產業？"),
                ("h2", "七、簡答題：共8分"),
                ("q", "28. 請用「原因 → 經過 → 影響」說明臺灣成為科技島的過程。4分。"),
                ("q", "29. 請舉一個臺灣多元文化的例子，並說明我們應該如何面對不同文化。4分。"),
                ("h2", "答案與解析摘要"),
                ("p", "選擇題：1B 2B 3A 4C 5B 6A 7B 8A 9D 10B。配對題：11B 12A 13C 14D。時序題：19為 B→C→A→D。"),
                ("p", "圖表與簡答題重點：15 產業由農業轉向工業，再發展高科技。16 工廠增加帶來就業機會，人口移入都市。17 可答人才、研發、資金、科學園區、政府政策。18 可答空氣或水汙染、土地開發、資源消耗。"),
                ("p", "情境題解析：21 地形與山脈阻隔仍會影響東西交通；交通改善可提升就醫、就學、工作、觀光與運輸便利。22 工業區能帶動就業與經濟，也須用環評、汙水處理與居民參與降低環境衝擊。23 訪問屬第一手資料，網路資料要查來源可信度。24 第二位同學較符合公民參與，因為提出具體且可行的公共行動。"),
                ("p", "閱讀與簡答解析：25 北部交通便利、工商業發達。26 山脈阻隔或地形因素。27 離島生活、物資、觀光與產業都受海洋交通影響。28 答題需含產業升級需求、科學園區與人才研發、半導體等產業影響。29 可舉米食、音樂、節慶、原住民族或新住民文化，態度須包含尊重、理解、欣賞。"),
            ],
        },
    ],
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "A7B0BA")
        borders.append(tag)
    tc_pr.append(borders)


def add_docx_para(doc, text, style_name=None, size=10.5, bold=False, color=None, space_after=2):
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.12
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "PingFang TC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_docx_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if r == 0 and any(x in rows[0][0] for x in ["時期"]):
                set_cell_shading(cell, "E8EEF5")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.name = "PingFang TC"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
            run.font.size = Pt(9.5)
            if r == 0 and any(x in rows[0][0] for x in ["時期"]):
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Cm(25)
    section.page_height = Cm(35.3)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    style = doc.styles["Normal"]
    style.font.name = "PingFang TC"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
    style.font.size = Pt(10.5)

    for page_index, section_data in enumerate(EXAM["sections"]):
        if page_index:
            doc.add_page_break()
        if page_index == 0:
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.space_after = Pt(1)
            run = title.add_run(EXAM["title"])
            run.font.name = "PingFang TC"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("12355B")
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.paragraph_format.space_after = Pt(6)
            r = sub.add_run(EXAM["subtitle"] + "　姓名：＿＿＿＿　班級：＿＿　座號：＿＿")
            r.font.name = "PingFang TC"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang TC")
            r.font.size = Pt(10)
        add_docx_para(doc, section_data["heading"], size=12, bold=True, color="12355B", space_after=4)
        for kind, content in section_data["blocks"]:
            if kind == "h2":
                add_docx_para(doc, content, size=10.5, bold=True, color="0B5F59", space_after=2)
            elif kind == "q":
                add_docx_para(doc, content, size=9.6, space_after=3)
            elif kind == "p":
                add_docx_para(doc, content, size=9.7, space_after=3)
            elif kind == "table":
                add_docx_table(doc, content)
    doc.save(DOCX_OUT)


def register_pdf_font():
    font_candidates = [
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for path in font_candidates:
        if Path(path).exists():
            pdfmetrics.registerFont(TTFont("CJK", path))
            return "CJK"
    return "Helvetica"


def pdf_para(text, style):
    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )
    return Paragraph(escaped, style)


def build_pdf():
    font = register_pdf_font()
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "base",
        parent=styles["Normal"],
        fontName=font,
        fontSize=9.4,
        leading=12.2,
        spaceAfter=4,
        alignment=TA_LEFT,
    )
    title = ParagraphStyle("title", parent=base, fontSize=17, leading=22, alignment=TA_CENTER, textColor=colors.HexColor("#12355B"), spaceAfter=2)
    subtitle = ParagraphStyle("subtitle", parent=base, fontSize=9.5, leading=12, alignment=TA_CENTER, textColor=colors.HexColor("#4B5563"), spaceAfter=8)
    h1 = ParagraphStyle("h1", parent=base, fontSize=11.8, leading=15, textColor=colors.HexColor("#12355B"), spaceBefore=2, spaceAfter=4)
    h2 = ParagraphStyle("h2", parent=base, fontSize=10.2, leading=13, textColor=colors.HexColor("#0B5F59"), spaceBefore=3, spaceAfter=2)

    doc = BaseDocTemplate(str(PDF_OUT), pagesize=B4, leftMargin=1.25 * cm, rightMargin=1.25 * cm, topMargin=1.25 * cm, bottomMargin=1.15 * cm)
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="B4", frames=[frame])])
    story = []
    for page_index, section_data in enumerate(EXAM["sections"]):
        if page_index:
            story.append(PageBreak())
        if page_index == 0:
            story.append(pdf_para(EXAM["title"], title))
            story.append(pdf_para(EXAM["subtitle"] + "　姓名：＿＿＿＿　班級：＿＿　座號：＿＿", subtitle))
        story.append(pdf_para(section_data["heading"], h1))
        for kind, content in section_data["blocks"]:
            if kind == "h2":
                story.append(pdf_para(content, h2))
            elif kind in ("q", "p"):
                story.append(pdf_para(content, base))
            elif kind == "table":
                data = [[pdf_para(str(x), base) for x in row] for row in content]
                col_count = len(content[0])
                widths = [doc.width / col_count] * col_count
                table = Table(data, colWidths=widths, hAlign="CENTER")
                style_cmds = [
                    ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#A7B0BA")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
                if "時期" in content[0][0]:
                    style_cmds.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")))
                table.setStyle(TableStyle(style_cmds))
                story.append(table)
                story.append(Spacer(1, 4))
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)
